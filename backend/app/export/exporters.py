import io
import markdown2
from docx import Document
from docx.shared import Pt


def to_markdown_bytes(md_text: str) -> bytes:
    return md_text.encode("utf-8")


def to_html_bytes(md_text: str, title: str = "Content") -> bytes:
    body = markdown2.markdown(md_text, extras=["fenced-code-blocks", "tables", "header-ids"])
    html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>{title}</title>
<style>
body {{ font-family: -apple-system, Segoe UI, Roboto, sans-serif; max-width: 800px; margin: 40px auto; line-height: 1.6; color:#1a1a1a; padding: 0 20px;}}
h1,h2,h3 {{ color: #111; }}
code {{ background:#f4f4f4; padding:2px 6px; border-radius:4px; }}
</style></head><body>{body}</body></html>"""
    return html.encode("utf-8")


def to_docx_bytes(md_text: str, title: str = "Content") -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    for line in md_text.split("\n"):
        line = line.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf_bytes(md_text: str, title: str = "Content") -> bytes:
    from weasyprint import HTML
    html_bytes = to_html_bytes(md_text, title)
    return HTML(string=html_bytes.decode("utf-8")).write_pdf()
